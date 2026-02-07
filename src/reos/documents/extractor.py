"""Text extraction from various document formats.

Supports: PDF, DOCX, TXT, MD, CSV, XLSX, XLS
"""

from __future__ import annotations

import csv
import logging
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


class DocumentExtractionError(Exception):
    """Error during document text extraction."""

    def __init__(self, message: str, file_path: Path | None = None) -> None:
        super().__init__(message)
        self.file_path = file_path


def extract_text(file_path: Path) -> tuple[str, dict[str, Any]]:
    """Extract text and metadata from a document.

    Args:
        file_path: Path to the document file.

    Returns:
        Tuple of (text_content, metadata_dict).
        Metadata includes: file_type, page_count (for PDFs), etc.

    Raises:
        DocumentExtractionError: If extraction fails.
    """
    if not file_path.exists():
        raise DocumentExtractionError(f"File not found: {file_path}", file_path)

    ext = file_path.suffix.lower()
    extractors = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".doc": _extract_docx,  # Try docx parser, may fail for old .doc
        ".txt": _extract_text_file,
        ".md": _extract_text_file,
        ".csv": _extract_csv,
        ".xlsx": _extract_excel,
        ".xls": _extract_excel,
    }

    extractor = extractors.get(ext)
    if not extractor:
        raise DocumentExtractionError(
            f"Unsupported file type: {ext}. Supported: {list(extractors.keys())}",
            file_path,
        )

    try:
        return extractor(file_path)
    except DocumentExtractionError:
        raise
    except Exception as exc:
        logger.exception("Failed to extract text from %s", file_path)
        raise DocumentExtractionError(
            f"Extraction failed: {exc}",
            file_path,
        ) from exc


def _extract_pdf(file_path: Path) -> tuple[str, dict[str, Any]]:
    """Extract text from PDF using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise DocumentExtractionError(
            "pypdf not installed. Install with: pip install pypdf",
            file_path,
        )

    reader = PdfReader(file_path)
    pages_text: list[str] = []
    page_boundaries: list[int] = []  # Character offsets for page starts

    current_offset = 0
    for page in reader.pages:
        page_boundaries.append(current_offset)
        text = page.extract_text() or ""
        pages_text.append(text)
        current_offset += len(text) + 1  # +1 for newline

    full_text = "\n".join(pages_text)

    metadata = {
        "file_type": "pdf",
        "page_count": len(reader.pages),
        "page_boundaries": page_boundaries,
    }

    # Extract PDF metadata if available
    if reader.metadata:
        if reader.metadata.title:
            metadata["title"] = reader.metadata.title
        if reader.metadata.author:
            metadata["author"] = reader.metadata.author

    return full_text, metadata


def _extract_docx(file_path: Path) -> tuple[str, dict[str, Any]]:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise DocumentExtractionError(
            "python-docx not installed. Install with: pip install python-docx",
            file_path,
        )

    doc = Document(file_path)
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n\n".join(paragraphs)

    metadata = {
        "file_type": "docx",
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
    }

    # Extract core properties if available
    try:
        if doc.core_properties.title:
            metadata["title"] = doc.core_properties.title
        if doc.core_properties.author:
            metadata["author"] = doc.core_properties.author
    except (AttributeError, KeyError) as e:
        logger.debug("Could not extract DOCX core properties: %s", e)

    return full_text, metadata


def _extract_text_file(file_path: Path) -> tuple[str, dict[str, Any]]:
    """Extract text from plain text or markdown files."""
    # Try common encodings
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]

    for encoding in encodings:
        try:
            text = file_path.read_text(encoding=encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise DocumentExtractionError(
            f"Could not decode file with encodings: {encodings}",
            file_path,
        )

    ext = file_path.suffix.lower()
    metadata = {
        "file_type": "markdown" if ext == ".md" else "text",
        "encoding": encoding,
        "line_count": text.count("\n") + 1,
    }

    return text, metadata


def _extract_csv(file_path: Path) -> tuple[str, dict[str, Any]]:
    """Extract text from CSV files."""
    # Try to detect encoding
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]

    for encoding in encodings:
        try:
            with open(file_path, newline="", encoding=encoding) as f:
                # Try to sniff the dialect
                sample = f.read(8192)
                f.seek(0)

                try:
                    dialect = csv.Sniffer().sniff(sample)
                except csv.Error:
                    dialect = csv.excel

                reader = csv.reader(f, dialect)
                rows = list(reader)
                break
        except UnicodeDecodeError:
            continue
    else:
        raise DocumentExtractionError(
            f"Could not decode CSV with encodings: {encodings}",
            file_path,
        )

    if not rows:
        return "", {"file_type": "csv", "row_count": 0, "column_count": 0}

    # Convert to markdown table format
    headers = rows[0] if rows else []
    lines: list[str] = []

    # Header row
    if headers:
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")

    # Data rows
    for row in rows[1:]:
        # Pad row to match header length
        padded = row + [""] * (len(headers) - len(row)) if len(row) < len(headers) else row
        lines.append("| " + " | ".join(padded[: len(headers)]) + " |")

    full_text = "\n".join(lines)

    metadata = {
        "file_type": "csv",
        "row_count": len(rows),
        "column_count": len(headers),
        "headers": headers,
    }

    return full_text, metadata


def _extract_excel(file_path: Path) -> tuple[str, dict[str, Any]]:
    """Extract text from Excel files using openpyxl."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise DocumentExtractionError(
            "openpyxl not installed. Install with: pip install openpyxl",
            file_path,
        )

    wb = load_workbook(file_path, read_only=True, data_only=True)
    sheets_text: list[str] = []
    total_rows = 0
    total_cols = 0

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        rows = list(sheet.iter_rows(values_only=True))

        if not rows:
            continue

        # Find actual data bounds (skip empty rows/cols)
        non_empty_rows = [row for row in rows if any(cell is not None for cell in row)]

        if not non_empty_rows:
            continue

        # Determine column count from first non-empty row
        col_count = len(non_empty_rows[0])
        total_cols = max(total_cols, col_count)
        total_rows += len(non_empty_rows)

        # Format as markdown table
        lines = [f"## Sheet: {sheet_name}\n"]

        # Header row
        headers = [str(cell) if cell is not None else "" for cell in non_empty_rows[0]]
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")

        # Data rows
        for row in non_empty_rows[1:]:
            cells = [str(cell) if cell is not None else "" for cell in row]
            # Pad to header length
            cells = (
                cells + [""] * (len(headers) - len(cells)) if len(cells) < len(headers) else cells
            )
            lines.append("| " + " | ".join(cells[: len(headers)]) + " |")

        sheets_text.append("\n".join(lines))

    wb.close()

    full_text = "\n\n".join(sheets_text)

    metadata = {
        "file_type": "excel",
        "sheet_count": len(wb.sheetnames),
        "sheet_names": wb.sheetnames,
        "total_rows": total_rows,
        "total_columns": total_cols,
    }

    return full_text, metadata
